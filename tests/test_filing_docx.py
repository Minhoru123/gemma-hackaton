from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from app import filing_docx

CAPTION = {"court": "IN THE UNITED STATES DISTRICT COURT FOR THE DISTRICT OF UTAH",
           "plaintiff": "David Jakeman", "defendant": "Cox",
           "case_no": "2:25-cv-00102", "judge": "Kimball",
           "title": "MOTION TO DISMISS"}


def test_build_produces_styled_docx_with_highlighted_markers(tmp_path):
    out = str(tmp_path / "filing.docx")
    filing_docx.build("First paragraph.\n\nDeadline is [[CONFIRM DATE]] today.",
                      CAPTION, out)
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert CAPTION["court"] in texts
    assert CAPTION["title"] in texts
    assert any("First paragraph." in t for t in texts)
    marker_runs = [r for p in doc.paragraphs for r in p.runs
                   if r.text == "[[CONFIRM DATE]]"]
    assert marker_runs
    assert marker_runs[0].bold is True
    assert marker_runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW
    body = [r for p in doc.paragraphs for r in p.runs if "First paragraph" in r.text]
    assert body[0].font.name == "Times New Roman"
