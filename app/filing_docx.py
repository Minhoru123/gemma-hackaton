"""Config-driven .docx filing builder: two-column caption, Times New Roman
13pt double-spaced body, unresolved [[MARKERS]] rendered bold-on-yellow so
they cannot slip into a filed document unnoticed."""

import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

_MARKER_SPLIT_RE = re.compile(r"(\[\[[^\]]+\]\])")

FONT = "Times New Roman"
SIZE = Pt(13)


def _style_run(run, bold=False, highlight=False):
    run.font.name = FONT
    run.font.size = SIZE
    run.bold = bold
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_body_paragraph(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    for part in _MARKER_SPLIT_RE.split(text):
        if not part:
            continue
        is_marker = part.startswith("[[")
        _style_run(p.add_run(part), bold=is_marker, highlight=is_marker)
    return p


def _add_centered(doc, text: str, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(p.add_run(text), bold=bold)
    return p


def build(draft_text: str, caption: dict, out_path: str) -> None:
    """caption keys: court, plaintiff, defendant, case_no, judge, title."""
    doc = Document()

    _add_centered(doc, caption.get("court", ""), bold=True)

    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left_text = (f"{caption.get('plaintiff', '')},\n    Plaintiff,\nv.\n"
                 f"{caption.get('defendant', '')},\n    Defendant.")
    right_lines = [f"Case No. {caption.get('case_no', '')}"]
    if caption.get("judge"):
        right_lines.append(f"Judge {caption['judge']}")
    for cell, text in ((left, left_text), (right, "\n".join(right_lines))):
        cell.paragraphs[0].text = ""
        for i, line in enumerate(text.split("\n")):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            _style_run(p.add_run(line))

    _add_centered(doc, caption.get("title", ""), bold=True)

    for para in draft_text.split("\n\n"):
        para = " ".join(para.split())
        if para:
            _add_body_paragraph(doc, para)

    doc.save(out_path)
